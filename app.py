import os
import io
import json
import pandas as pd
from datetime import datetime
from flask import Flask, render_template, request, jsonify, send_file
from google import genai
from google.genai import types
from dotenv import load_dotenv

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

load_dotenv()

app = Flask(__name__)

client = genai.Client(api_key=os.environ.get("GEMINI_API_KEY"))
GEMINI_MODEL = "gemini-3.1-flash-lite"

df = pd.read_csv('curriculum.csv').fillna('')


@app.template_filter('render_list')
def render_list_filter(text):
    """Convert plain-text numbered/bulleted content to semantic HTML for display."""
    import re, html
    if not text:
        return ''
    lines = text.strip().split('\n')
    out   = []
    mode  = None   # 'ol', 'ul', or None

    def close_list():
        nonlocal mode
        if mode == 'ol': out.append('</ol>')
        if mode == 'ul': out.append('</ul>')
        mode = None

    for raw in lines:
        line = raw.strip()
        if not line:
            close_list()
            out.append('<br>')
            continue

        num  = re.match(r'^(\d+)\.\s+(.+)$', line)
        bull = re.match(r'^[•\-\*]\s+(.+)$', line)
        # Bold label lines like "Assessment:", "Closure:", "Pre-reading:"
        label = re.match(r'^(Assessment:|Closure:|Pre-reading:|Post-reading[/\w]*:?)\s*(.*)', line, re.I)

        if num:
            if mode != 'ol': close_list(); out.append('<ol>'); mode = 'ol'
            out.append(f'<li>{html.escape(num.group(2))}</li>')
        elif bull:
            if mode != 'ul': close_list(); out.append('<ul>'); mode = 'ul'
            out.append(f'<li>{html.escape(bull.group(1))}</li>')
        elif label:
            close_list()
            rest = html.escape(label.group(2))
            out.append(f'<p><strong>{html.escape(label.group(1))}</strong>{" " + rest if rest else ""}</p>')
        else:
            close_list()
            out.append(f'<p>{html.escape(line)}</p>')

    close_list()
    return ''.join(out)


# ── Routes ──────────────────────────────────────────────────────────────────

@app.route('/')
def index():
    levels = sorted(df['Class_Level'].unique().tolist())
    today  = datetime.today().strftime('%Y-%m-%d')
    return render_template('index.html', levels=levels, today=today)


@app.route('/api/substrands')
def get_all_substrands():
    return jsonify(sorted(df['Sub_Strand_Number_and_Name'].unique().tolist()))


@app.route('/api/substrands/<level>')
def get_substrands_by_level(level):
    filtered = df[df['Class_Level'] == level]
    return jsonify(sorted(filtered['Sub_Strand_Number_and_Name'].unique().tolist()))


@app.route('/generate', methods=['POST'])
def generate():
    week_starting = request.form.get('week_starting', '')
    week_ending   = request.form.get('week_ending', '')
    class_size    = request.form.get('class_size', '35')

    try:
        week_label = datetime.strptime(week_starting, '%Y-%m-%d').strftime('Week of %B %d, %Y')
    except ValueError:
        week_label = week_starting

    sub_strands      = request.form.getlist('sub_strand[]')
    dates            = request.form.getlist('date[]')
    durations        = request.form.getlist('duration[]')
    custom_durations = request.form.getlist('custom_duration[]')

    resolved_durations = []
    for i, d in enumerate(durations):
        if d == 'custom' and i < len(custom_durations) and custom_durations[i].strip():
            resolved_durations.append(f"{custom_durations[i].strip()} mins")
        else:
            resolved_durations.append(d)

    strand_counts   = {}
    for ss in sub_strands:
        strand_counts[ss] = strand_counts.get(ss, 0) + 1

    current_tracker = {}
    lessons_data    = []

    for i, sub_strand in enumerate(sub_strands):
        current_tracker[sub_strand] = current_tracker.get(sub_strand, 0) + 1
        lesson_num    = current_tracker[sub_strand]
        total_lessons = strand_counts[sub_strand]

        rows = df[df['Sub_Strand_Number_and_Name'] == sub_strand]
        if rows.empty:
            continue
        row = rows.iloc[0]

        duration_str = resolved_durations[i] if i < len(resolved_durations) else '60 mins'
        p1_time, p2_time, p3_time = split_phase_times(duration_str)

        ai_draft = generate_ai_drafts(
            sub_strand       = sub_strand,
            indicator        = row['Indicator'],
            exemplars        = row['Phase_2_Main_Exemplers'],
            competencies     = row['Core_Competencies_Values'],
            content_standard = row['Content_Standard'],
            class_level      = row['Class_Level'],
            lesson_num       = lesson_num,
            total_lessons    = total_lessons,
        )

        lessons_data.append({
            "Subject":                    row['Subject'],
            "Class_Level":                row['Class_Level'],
            "Class_Size":                 class_size,
            "Week_Ending":                week_ending,
            "Date":                       dates[i] if i < len(dates) else '',
            "Duration":                   duration_str,
            "Phase_1_Time":               p1_time,
            "Phase_2_Time":               p2_time,
            "Phase_3_Time":               p3_time,
            "Lesson_Number":              f"Lesson {lesson_num} of {total_lessons}",
            "Strand_Number_and_Name":     row['Strand_Number_and_Name'],
            "Sub_Strand_Number_and_Name": sub_strand,
            "Content_Standard":           row['Content_Standard'],
            "Indicator":                  row['Indicator'],
            "Performance_Indicator":      ai_draft.get("performance_indicator", ""),
            "Core_Competencies_Values":   row['Core_Competencies_Values'],
            "Key_Words":                  ", ".join(ai_draft.get("key_words", [])),
            "Resources_TLMs":             ai_draft.get("resources", ""),
            "Phase_1_Starter":            normalize_list_text(ai_draft.get("phase_1", "")),
            "Phase_2_Main":               normalize_list_text(ai_draft.get("phase_2", row['Phase_2_Main_Exemplers'])),
            "Assessment_Task":            ai_draft.get("assessment", ""),
            "Phase_3_Plenary_Reflections": normalize_list_text(ai_draft.get("phase_3", "")),
        })

    return render_template('lesson_plan.html', lessons=lessons_data)


@app.route('/download-docx', methods=['POST'])
def download_docx():
    lessons_data = request.json.get('lessons', [])
    buf = build_docx(lessons_data)
    return send_file(
        buf,
        as_attachment=True,
        download_name='Weekly_Lesson_Plans.docx',
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    )


# ── Helpers ─────────────────────────────────────────────────────────────────

def normalize_list_text(text):
    """Ensure each numbered item and key labels start on their own line."""
    import re
    if not text:
        return text
    # Insert newline before "N. " items that follow sentence-ending punctuation mid-text
    text = re.sub(r'([.!?])\s+(\d+\.\s)', r'\1\n\2', text)
    # Insert blank line before "Assessment:" / "Closure:" labels when not already at line start
    text = re.sub(r'(?<!\n)(Assessment:|Closure:)', r'\n\n\1', text)
    # Collapse 3+ newlines back to 2
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()


def split_phase_times(duration_str):
    """Return (phase1, phase2, phase3) minute strings from total duration."""
    import re
    m = re.search(r'(\d+)', str(duration_str))
    if not m:
        return '5 minutes', '20 minutes', '5 minutes'
    total = int(m.group(1))
    p1 = max(5, round(total / 6))
    p3 = max(5, round(total / 6))
    p2 = total - p1 - p3
    return f'{p1} minutes', f'{p2} minutes', f'{p3} minutes'


# ── AI generation ────────────────────────────────────────────────────────────

def generate_ai_drafts(sub_strand, indicator, exemplars, competencies,
                       content_standard, class_level, lesson_num=1, total_lessons=1):

    if total_lessons == 1:
        focus = "This is a single standalone lesson. Cover the topic fully and comprehensively."
    elif lesson_num == 1:
        focus = (
            f"LESSON {lesson_num} OF {total_lessons}. "
            "Focus on INTRODUCTION: activate prior knowledge, introduce vocabulary, "
            "use the FIRST exemplar activities for Phase 2 — keep it foundational."
        )
    elif lesson_num == total_lessons:
        focus = (
            f"LESSON {lesson_num} OF {total_lessons}. "
            "Students already had the introduction. "
            "Focus on CONSOLIDATION: deeper group tasks, extended application, "
            "LATER exemplar activities, substantial assessment."
        )
    else:
        focus = (
            f"LESSON {lesson_num} OF {total_lessons}. "
            "Build on Lesson 1. Focus on DEVELOPMENT: collaborative tasks, "
            "guided application, mid-level exemplar activities."
        )

    prompt = f"""You are an expert Ghanaian curriculum developer writing one NaCCA lesson plan.

Curriculum context:
- Class Level: {class_level}
- Sub-strand: {sub_strand}
- Content Standard: {content_standard}
- Indicator: {indicator}
- Core Competencies: {competencies}
- Curriculum Exemplars (use heavily for Phase 2): {exemplars}

Lesson context: {focus}

Return ONLY a valid JSON object — no markdown, no code fences — with these exact keys:
- "performance_indicator": string — one sentence starting with "The learner should be able to..." derived from the Indicator above
- "key_words": list of 4-6 key vocabulary words or phrases from this lesson
- "resources": string — specific Teaching/Learning Materials (TLMs) needed for Phase 2
- "phase_1": string — Starter activity (preparing the brain for learning). Bullet-point steps.
- "phase_2": string — Main teaching phase adapted from the exemplars with numbered steps and teacher-learner interactions. End with a short paragraph starting "Assessment:" describing an AfL task.
- "phase_3": string — Plenary/Reflections (Learner and teacher). Ask learners to reflect and connect to personal experience."""

    try:
        response = client.models.generate_content(
            model=GEMINI_MODEL,
            contents=prompt,
            config=types.GenerateContentConfig(temperature=0.3, max_output_tokens=1500),
        )
        content = response.text.strip()
        if content.startswith("```"):
            parts = content.split("```")
            content = parts[1] if len(parts) > 1 else content
            if content.startswith("json"):
                content = content[4:]
        return json.loads(content.strip())
    except Exception as e:
        print(f"AI Generation Error: {e}")
        return {
            "performance_indicator": f"The learner should be able to demonstrate understanding of {sub_strand}.",
            "key_words": [],
            "resources": "Textbooks, chalkboard, markers, flashcards",
            "phase_1": "Review the previous lesson. Ask learners 2–3 questions about what they learned.",
            "phase_2": exemplars + "\n\nAssessment: Ask learners oral questions to check understanding.",
            "phase_3": "Ask learners to summarise the key points of the lesson in their own words.",
        }


# ── Word document builder ────────────────────────────────────────────────────

def _shade_cell(cell, hex_fill):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_fill)
    tcPr.append(shd)


def _write_cell(cell, segments, font_size=9, align=None):
    """
    segments: list of (text, bold) tuples.
    Clears any existing content and writes fresh.
    """
    cell.text = ''
    para = cell.paragraphs[0]
    if align is not None:
        para.alignment = align
    for text, bold in segments:
        run = para.add_run(text)
        run.bold = bold
        run.font.size = Pt(font_size)


def build_docx(lessons):
    doc = Document()
    for sec in doc.sections:
        sec.top_margin    = Cm(1.5)
        sec.bottom_margin = Cm(1.5)
        sec.left_margin   = Cm(1.8)
        sec.right_margin  = Cm(1.8)

    GH_GREEN = RGBColor(0, 107, 63)

    for idx, L in enumerate(lessons):
        if idx > 0:
            doc.add_page_break()

        # ── Heading ──
        tp = doc.add_paragraph()
        tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = tp.add_run(f"GHANAIAN LANGUAGE  —  {L.get('Lesson_Number','')}")
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = GH_GREEN

        # ── Table 1: header info (3 rows × 2 cols) ──
        t1 = doc.add_table(rows=3, cols=2)
        t1.style = 'Table Grid'
        _write_cell(t1.rows[0].cells[0],
            [('Date: ', True), (L.get('Date',''), False),
             ('   Week Ending: ', True), (L.get('Week_Ending',''), False)])
        _write_cell(t1.rows[0].cells[1],
            [('Subject: ', True), (L.get('Subject',''), False)])
        _write_cell(t1.rows[1].cells[0],
            [('Duration: ', True), (L.get('Duration',''), False)])
        _write_cell(t1.rows[1].cells[1],
            [('Strand: ', True), (L.get('Strand_Number_and_Name',''), False)])
        _write_cell(t1.rows[2].cells[0],
            [('Class: ', True), (L.get('Class_Level',''), False),
             ('   Class size: ', True), (L.get('Class_Size','35'), False)])
        _write_cell(t1.rows[2].cells[1],
            [('Sub strand: ', True), (L.get('Sub_Strand_Number_and_Name',''), False)])

        # ── Table 2: standards (3 rows) ──
        t2 = doc.add_table(rows=4, cols=3)
        t2.style = 'Table Grid'
        # Row 1: Content Standard | Indicator | Lesson X of Y
        _write_cell(t2.rows[0].cells[0],
            [('Content Standard: ', True), (L.get('Content_Standard',''), False)])
        _write_cell(t2.rows[0].cells[1],
            [('Indicator: ', True), (L.get('Indicator',''), False)])
        _write_cell(t2.rows[0].cells[2],
            [(L.get('Lesson_Number',''), True)],
            align=WD_ALIGN_PARAGRAPH.CENTER)
        # Row 2: Performance Indicator (spans 2) | Core Competencies
        t2.rows[1].cells[0].merge(t2.rows[1].cells[1])
        _write_cell(t2.rows[1].cells[0],
            [('Performance Indicator: ', True), (L.get('Performance_Indicator',''), False)])
        _write_cell(t2.rows[1].cells[2],
            [('Core Competencies/Values: ', True), (L.get('Core_Competencies_Values',''), False)],
            font_size=8)
        # Row 3: Key words (full width)
        t2.rows[2].cells[0].merge(t2.rows[2].cells[2])
        _write_cell(t2.rows[2].cells[0],
            [('Key words: ', True), (L.get('Key_Words',''), False)])
        # Row 4: Phase table header
        for cell, txt in zip(t2.rows[3].cells,
                              ['Phase/Duration', 'Learner activities', 'Resources']):
            _write_cell(cell, [(txt, True)])
            _shade_cell(cell, 'D4EDDA')

        # ── Table 3: phase rows (3 rows × 3 cols) ──
        t3 = doc.add_table(rows=3, cols=3)
        t3.style = 'Table Grid'

        p1_label = f"Phase 1:\nStarter\n{L.get('Phase_1_Time','5 minutes')}"
        p2_label = f"Phase 2:\nMain\n{L.get('Phase_2_Time','20 minutes')}"
        p3_label = f"Phase 3:\nPlenary/\nReflections\n{L.get('Phase_3_Time','5 minutes')}"

        phase_content = [
            (p1_label, [(L.get('Phase_1_Starter',''), False)],            ''),
            (p2_label, [(L.get('Phase_2_Main',''), False)],               L.get('Resources_TLMs','')),
            (p3_label, [(L.get('Phase_3_Plenary_Reflections',''), False)], ''),
        ]

        for r_idx, (label, content_segs, resources) in enumerate(phase_content):
            row = t3.rows[r_idx]
            _write_cell(row.cells[0], [(label, False)], align=WD_ALIGN_PARAGRAPH.LEFT)
            _shade_cell(row.cells[0], 'F1F8F4')
            _write_cell(row.cells[1], content_segs)
            _write_cell(row.cells[2], [(resources, False)])

        # Approximate column widths
        for tbl in [t2, t3]:
            for col, w in zip(tbl.columns, [Cm(3.8), Cm(10), Cm(4.2)]):
                for cell in col.cells:
                    cell.width = w

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


if __name__ == '__main__':
    app.run(debug=True)
